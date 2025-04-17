from docx import Document
from docx.shared import Inches
import numpy as np
import os
import pandas as pd
import pickle
from sklearn.model_selection import train_test_split
from radiomics_processes.extract_features.radiomics_extractor import RadiomicFeatureExtractor
from radiomics_processes.utils import verify_dataset_intergrity
from radiomics_processes.preprocessing.feature_preprocessing import run_preprocessing
from radiomics_processes.select_features.feature_selector import FeatureSelector
from radiomics_processes.classify_process.classification_base import feature_classification_process
from radiomics_processes.metrics_reports.report_results import ResultReportor
from sklearn.utils import shuffle


class Trainer(object):
    def __init__(self, train_dir, patient_info, test_size, save_dir, modility_for_extraction,
                 radiomic_feature_types, image_filters, feature_preprocessing_method, feature_selection_methods,
                 classifier, random_seed):
        """
        feature extraction and selection and classification
        :param train_dir:
        :param patient_info:
        :param test_size:
        :param save_dir:
        :param type_tag:
        :param modility_for_extraction:
        # :param label_tag:
        :param radiomic_feature_types:
        :param image_filters:
        :param feature_preprocessing_method:
        :param feature_selection_methods:
        :param classifier:
        :param random_seed:
        """
        self.train_dir = train_dir
        self.patient_info = patient_info
        self.test_size = test_size
        self.save_dir = save_dir
        os.makedirs(save_dir, exist_ok=True)
        self.modility_for_extraction = modility_for_extraction
        # self.label_tag = label_tag

        self.radiomic_feature_types = radiomic_feature_types
        self.image_filters = image_filters
        self.feature_preprocessing_method = feature_preprocessing_method
        self.feature_selection_methods = feature_selection_methods
        self.classifier = classifier
        self.random_seed = random_seed
        self.document = Document()

    def analysis_data(self):
        """
        get the distribution of positive and negative in train, test class
        :return:
        """
        all_numbers = len(self.feature_data)
        num_feature = len(self.feature_data.columns) - 2  # the other two is case_id and label
        pos_all = np.sum(self.feature_data['label'].values)
        neg_all = all_numbers - pos_all
        train_number = len(self.train_set)
        test_number = len(self.test_set)
        train_pos = np.sum(self.train_set['label'].values)
        train_neg = train_number - train_pos
        test_pos = np.sum(self.test_set['label'].values)
        test_neg = test_number - test_pos
        data_report = (
            ('data', all_numbers, pos_all, neg_all),
            ('train', train_number, train_pos, train_neg),
            ('test', test_number, test_pos, test_neg)
        )
        # write the table to doc
        table = self.document.add_table(rows=1, cols=4, style='Table Grid')
        row = table.rows[0].cells
        row[0].text = 'type'
        row[1].text = 'sum'
        row[2].text = 'positive'
        row[3].text = 'negative'
        for type, sum_, pos, neg in data_report:
            row = table.add_row().cells
            row[0].text = type
            row[1].text = str(sum_)
            row[2].text = str(pos)
            row[3].text = str(neg)

    def run_process(self):
        print('1.verify the data intergrity!')
        verify_dataset_intergrity(self.train_dir, self.patient_info, self.modility_for_extraction)
        print('2.extract radiomic features')
        if os.path.exists(os.path.join(self.save_dir, 'radiomic_features.xlsx')):
            self.feature_data_original = pd.read_excel(os.path.join(self.save_dir, 'radiomic_features.xlsx'))
            print('the number of feature extracted is:', self.feature_data_original.shape[1])
        else:
            extractor = RadiomicFeatureExtractor(self.train_dir, self.patient_info, self.save_dir,
                                                 self.modility_for_extraction, self.radiomic_feature_types,
                                                 self.image_filters)

            self.feature_data_original = extractor.run_extraction()

        print('3.feature preprocessing')
        if self.feature_preprocessing_method == 0:
            feature_preprocess_tag = 'Standardization'
        else:
            feature_preprocess_tag = 'min-max rescaling'

        if os.path.exists(os.path.join(self.save_dir, 'radiomics_features_after_preprocessing.xlsx')):
            self.feature_data = pd.read_excel(
                os.path.join(self.save_dir, 'radiomics_features_after_preprocessing.xlsx'))
        else:
            self.feature_data = run_preprocessing(self.feature_data_original, self.feature_preprocessing_method)
            self.feature_data.to_excel(os.path.join(self.save_dir, 'radiomics_features_after_preprocessing.xlsx'),
                                       index=False)

        print('4.split the data to train set and test set')
        labels = self.feature_data['label'].values
        self.train_set, self.test_set = train_test_split(self.feature_data, test_size=self.test_size,
                                                         random_state=self.random_seed, stratify=labels)
        # print(self.train_set.shape)
        self.document.add_heading('Summary Report', 0)
        self.document.add_heading('I. data information', level=1)
        self.document.add_heading('1. data location', level=2)
        self.document.add_paragraph(
            'data set location: {}'.format(self.train_dir), style='List Number'
        )
        self.document.add_heading('2. random seed for split is {}'.format(self.random_seed), level=2)
        self.document.add_heading('3. test size is {}'.format(self.test_size), level=2)
        self.analysis_data()  # analysis the distribution of positive and negative

        self.document.add_heading('II. method part', level=1)
        self.document.add_heading('1. the parameters for feature extraction')
        self.document.add_paragraph('the modalities used here is {}'.format(self.modility_for_extraction))
        self.document.add_paragraph('the feature types here is {}'.format(self.radiomic_feature_types))
        self.document.add_paragraph('the image filters used here is {}'.format(self.image_filters))
        self.document.add_paragraph(
            'the number of feature extracted is {}'.format(self.feature_data_original.shape[1] - 2))
        self.document.add_paragraph(
            'the number of feature after remove low variance is {}'.format(self.feature_data.shape[1] - 2))

        self.document.add_heading('2. the method for standardizing the feature is:{}'.format(feature_preprocess_tag),
                                  level=2)

        print('5.feature selection')
        selector = FeatureSelector(self.feature_selection_methods, self.save_dir)
        new_train_set, new_test_set = selector.run_selection(self.train_set, self.test_set)

        print('6.feature classification')
        train_probs, test_probs = feature_classification_process(self.classifier, new_train_set, new_test_set,
                                                                 self.save_dir, self.random_seed)
        analysis_processor = ResultReportor(test_probs, train_probs, self.save_dir, self.document)
        analysis_processor.run_analysis()


class TrainerWithFeatures(object):
    def __init__(self, feature_path, test_size, save_dir, feature_selection_methods, unbalanced, classifier,
                 random_seed):

        """
        feature selection and classification without extracting and preprocessing such as variance reduction and normalization
        :param train_dir:
        :param patient_info:
        :param test_size:
        :param save_dir:
        :param type_tag:
        :param modility_for_extraction:
        # :param label_tag:
        :param radiomic_feature_types:
        :param image_filters:
        :param feature_preprocessing_method:
        :param feature_selection_methods:
        :param classifier:
        :param random_seed:
        """
        self.feature_path = feature_path
        # self.unbalanced = unbalanced
        if unbalanced:
            self.class_weight = "balanced"
        else:
            self.class_weight = None

        self.test_size = test_size
        self.save_dir = save_dir
        os.makedirs(save_dir, exist_ok=True)

        # self.label_tag = label_tag

        self.feature_selection_methods = feature_selection_methods
        self.classifier = classifier
        self.random_seed = random_seed
        self.document = Document()

    def analysis_data(self):
        """
        get the distribution of positive and negative in train, test class
        :return:
        """
        all_numbers = len(self.feature_data)
        num_feature = len(self.feature_data.columns) - 2  # the other two is case_id and label
        pos_all = np.sum(self.feature_data['label'].values)
        neg_all = all_numbers - pos_all
        train_number = len(self.train_set)
        test_number = len(self.test_set)
        train_pos = np.sum(self.train_set['label'].values)
        train_neg = train_number - train_pos
        test_pos = np.sum(self.test_set['label'].values)
        test_neg = test_number - test_pos

        data_report = (
            ('data', all_numbers, pos_all, neg_all),
            ('train', train_number, train_pos, train_neg),
            ('test', test_number, test_pos, test_neg)
        )
        # write the table to doc
        table = self.document.add_table(rows=1, cols=4, style='Table Grid')
        row = table.rows[0].cells
        row[0].text = 'type'
        row[1].text = 'sum'
        row[2].text = 'positive'
        row[3].text = 'negative'
        for type, sum_, pos, neg in data_report:
            row = table.add_row().cells
            row[0].text = type
            row[1].text = str(sum_)
            row[2].text = str(pos)
            row[3].text = str(neg)

    def run_process(self):
        # print('1.verify the data intergrity!')
        # verify_dataset_intergrity(self.train_dir, self.patient_info, self.modility_for_extraction)
        # print('2.extract radiomic features')
        # if os.path.exists(os.path.join(self.save_dir, 'radiomic_features.xlsx')):
        #     self.feature_data_original = pd.read_excel(os.path.join(self.save_dir, 'radiomic_features.xlsx'))
        #     print('the number of feature extracted is:', self.feature_data_original.shape[1])
        # else:
        #     extractor = RadiomicFeatureExtractor(self.train_dir, self.patient_info, self.save_dir,
        #                                          self.modility_for_extraction, self.radiomic_feature_types,
        #                                          self.image_filters)
        #
        #     self.feature_data_original = extractor.run_extraction()

        # print('3.feature preprocessing')
        # if self.feature_preprocessing_method == 0:
        #     feature_preprocess_tag = 'Standardization'
        # else:
        #     feature_preprocess_tag = 'min-max rescaling'
        #
        try:
            self.feature_data = pd.read_csv(self.feature_path)
            self.feature_data = shuffle(self.feature_data)

        except FileNotFoundError:
            print(f"Sorry,the file {self.feature_path} does not exist.")

            # self.feature_data = run_preprocessing(self.feature_data_original, self.feature_preprocessing_method)
            # self.feature_data.to_excel(os.path.join(self.save_dir, 'radiomics_features_after_preprocessing.xlsx'),
            #                            index=False)

        print('1.split the data to train set and test set')
        labels = self.feature_data['label'].values
        self.train_set, self.test_set = train_test_split(self.feature_data, test_size=self.test_size, shuffle=True,
                                                         random_state=self.random_seed, stratify=labels)
        # with open('/home/lsq/dataset/Fu1ShenQian/radiomics_features/preprocessed_features/data_split.pkl', 'rb') as f:
        #     data_split = pickle.load(f)
        # self.train_set = self.feature_data.loc[self.feature_data['case_ids'].isin(data_split['train'])]
        # self.test_set = self.feature_data.loc[self.feature_data['case_ids'].isin(data_split['test'])]

        # print(self.train_set.shape)
        self.document.add_heading('Summary Report', 0)
        self.document.add_heading('I. data information', level=1)
        self.document.add_heading('1. data location', level=2)
        self.document.add_paragraph(
            'feature location: {}'.format(self.feature_path), style='List Number'
        )
        self.document.add_heading('2. random seed for split is {}'.format(self.random_seed), level=2)
        self.document.add_heading('3. test size is {}'.format(self.test_size), level=2)
        self.analysis_data()  # analysis the distribution of positive and negative

        self.document.add_heading('II. method part', level=1)
        # self.document.add_heading('1. the parameters for feature extraction')
        # self.document.add_paragraph('the modalities used here is {}'.format(self.modility_for_extraction))
        # self.document.add_paragraph('the feature types here is {}'.format(self.radiomic_feature_types))
        # self.document.add_paragraph('the image filters used here is {}'.format(self.image_filters))
        # self.document.add_paragraph(
        #     'the number of feature extracted is {}'.format(self.feature_data_original.shape[1] - 2))
        # self.document.add_paragraph(
        #     'the number of feature after remove low variance is {}'.format(self.feature_data.shape[1] - 2))

        # self.document.add_heading('2. the method for standardizing the feature is:{}'.format(feature_preprocess_tag),
        #                           level=2)

        print('2.feature selection')
        selector = FeatureSelector(self.feature_selection_methods, self.save_dir)
        new_train_set, new_test_set = selector.run_selection(self.train_set, self.test_set)

        print('3.feature classification')
        train_probs, test_probs = feature_classification_process(self.classifier, new_train_set, new_test_set,
                                                                 self.save_dir, self.class_weight, self.random_seed)
        analysis_processor = ResultReportor(test_probs, train_probs, self.save_dir, self.document)
        analysis_processor.run_analysis()
